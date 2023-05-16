import numpy as np
from sys import stdout
from openmm.app import *
from openmm import *
from openmm.unit import *
import statistics
import matplotlib.pyplot as plt
import docx
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Qt5Agg')

""" program to calculate the atom-center of mass distance and Rg from a pdb file. OA 20230122"""
myPdb = PDBFile('BoltornG4Dry30M.pdb')
#Time of the step in 30M run
time_30M = np.linspace(2*30000000/20, 2*30000000, 20)

positions_1 = myPdb.getPositions(frame=0)  #get the positions for the frame in nanometer
X_1 = [float(coord[0]/nanometer) for coord in positions_1]
Y_1 = [float(coord[1]/nanometer) for coord in positions_1]
Z_1 = [float(coord[2]/nanometer) for coord in positions_1]
noAtoms = len(X_1)#Assign of the number of atoms
#Assign of atoms
myAtoms = [atom for atom in myPdb.topology.atoms()]
# The elements are returned as <Element carbon>
myElements = [str(atom.element)[9:-1] for atom in myAtoms] #Assign of elemets names
myMasses =[]# atom masses
newX_1 = [] #The coordinates of the atoms with respect to center of mass
newY_1 = []
newZ_1 = []
for i in range(noAtoms):
    if myElements[i] == 'hydrogen':
        myMasses.append(1.00709)
    if myElements[i] == 'carbon':
        myMasses.append(12.011)
    if myElements[i] == 'oxygen':
        myMasses.append(15.999)

    if isinstance(X_1[i], (int, float)):
        newX_1.append(X_1[i] - X_1[0])

    if isinstance(Y_1[i], (int, float)):
        newY_1.append(Y_1[i] - Y_1[0])

    if isinstance(Z_1[i], (int, float)):
        newZ_1.append(Z_1[i] - Z_1[0])
#Total number of mass
totalMass = sum(myMasses)
#(atom-center of mass) distances
myDistances_1 = np.sqrt([(newX_1[i]**2) + (newY_1[i]**2)+ (newZ_1[i]**2) for i in range(noAtoms)])
mySquaredDistances_1 = [] 
for x in myDistances_1:
    mySquaredDistances_1.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_1 = [mySquaredDistances_1[i] * myMasses[i] for i in range(noAtoms)]
mySquare_1Sum = sum(myMassWeightedSqDists_1)
# Radius of gyration of the structure
radiusOfGyration_1 = np.sqrt(mySquare_1Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_1 = ', sum(myDistances_1)/(noAtoms), 'nm')   
print('Radius of Gyration_1 is:', radiusOfGyration_1, 'nm')
#(atom-center of mass) distance
print('distance_1(atom-center of mass) = ', myDistances_1)

positions_2 = myPdb.getPositions(frame=1)  #get the positions for the frame n nanometer
X_2 = [float(coord[0]/nanometer) for coord in positions_2]
Y_2 = [float(coord[1]/nanometer) for coord in positions_2]
Z_2 = [float(coord[2]/nanometer) for coord in positions_2]
newX_2 = [] #The coordinates of the atoms with respect to center of mass
newY_2 = []
newZ_2 = []
for i in range(noAtoms):

    if isinstance(X_2[i], (int, float)):
        newX_2.append(X_2[i] - X_2[0])

    if isinstance(Y_2[i], (int, float)):
        newY_2.append(Y_2[i] - Y_2[0])

    if isinstance(Z_2[i], (int, float)):
        newZ_2.append(Z_2[i] - Z_2[0])
#(atom-center of mass) distances
myDistances_2 = np.sqrt([(newX_2[i]**2) + (newY_2[i]**2)+ (newZ_2[i]**2) for i in range(noAtoms)])
mySquaredDistances_2 = [] 
for x in myDistances_2:
    mySquaredDistances_2.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_2 = [mySquaredDistances_2[i] * myMasses[i] for i in range(noAtoms)]
mySquare_2Sum = sum(myMassWeightedSqDists_2)
# Radius of gyration of the structure
radiusOfGyration_2 = np.sqrt(mySquare_2Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_2 = ', sum(myDistances_2)/(noAtoms), 'nm')   
print('Radius of Gyration_2 is:', radiusOfGyration_2, 'nm')
#(atom-center of mass) distances
print('distance_2(atom-center of mass) = ',myDistances_2)

positions_3 = myPdb.getPositions(frame=2) #get the positions for the frame in nanometer
X_3 = [float(coord[0]/nanometer) for coord in positions_3]
Y_3 = [float(coord[1]/nanometer) for coord in positions_3]
Z_3 = [float(coord[2]/nanometer) for coord in positions_3]
newX_3 = [] #The coordinates of the atoms with respect to center of mass
newY_3 = []
newZ_3 = []
for i in range(noAtoms):

    if isinstance(X_3[i], (int, float)):
        newX_3.append(X_3[i] - X_3[0])

    if isinstance(Y_3[i], (int, float)):
        newY_3.append(Y_3[i] - Y_3[0])

    if isinstance(Z_3[i], (int, float)):
        newZ_3.append(Z_3[i] - Z_3[0])
#(atom-center of mass) distances
myDistances_3 = np.sqrt([(newX_3[i]**2) + (newY_3[i]**2)+ (newZ_3[i]**2) for i in range(noAtoms)])
mySquaredDistances_3 = [] 
for x in myDistances_3:
    mySquaredDistances_3.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_3 = [mySquaredDistances_3[i] * myMasses[i] for i in range(noAtoms)]
mySquare_3Sum = sum(myMassWeightedSqDists_3)
# Radius of gyration of the structure
radiusOfGyration_3 = np.sqrt(mySquare_3Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_3 = ', sum(myDistances_3)/(noAtoms), 'nm')   
print('Radius of Gyration_3 is:', radiusOfGyration_3, 'nm')
#(atom-center of mass) distances
print('distance_3 = ', myDistances_3)


positions_4 = myPdb.getPositions(frame=3) #get the positions for the frame in nanometer
X_4 = [float(coord[0]/nanometer) for coord in positions_4]
Y_4 = [float(coord[1]/nanometer) for coord in positions_4]
Z_4 = [float(coord[2]/nanometer) for coord in positions_4]
newX_4 = [] #The coordinates of the atoms with respect to center of mass
newY_4 = []
newZ_4 = []
for i in range(noAtoms):

    if isinstance(X_4[i], (int, float)):
        newX_4.append(X_4[i] - X_4[0])

    if isinstance(Y_4[i], (int, float)):
        newY_4.append(Y_4[i] - Y_4[0])

    if isinstance(Z_4[i], (int, float)):
        newZ_4.append(Z_4[i] - Z_4[0])
#(atom-center of mass) distances
myDistances_4 = np.sqrt([(newX_4[i]**2) + (newY_4[i]**2)+ (newZ_4[i]**2) for i in range(noAtoms)])
mySquaredDistances_4 = [] 
for x in myDistances_4:
    mySquaredDistances_4.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_4 = [mySquaredDistances_4[i] * myMasses[i] for i in range(noAtoms)]
mySquare_4Sum = sum(myMassWeightedSqDists_4)
# Radius of gyration of the structure
radiusOfGyration_4 = np.sqrt(mySquare_4Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_4 = ', sum(myDistances_4)/(noAtoms), 'nm')   
print('Radius of Gyration_4 is:', radiusOfGyration_4, 'nm')
#(atom-center of mass) distances
print('distance_4 = ', myDistances_4)


positions_5 = myPdb.getPositions(frame=4) #get the positions for the frame in nanometer
X_5 = [float(coord[0]/nanometer) for coord in positions_5]
Y_5 = [float(coord[1]/nanometer) for coord in positions_5]
Z_5 = [float(coord[2]/nanometer) for coord in positions_5]
newX_5 = [] #The coordinates of the atoms with respect to center of mass
newY_5 = []
newZ_5 = []
for i in range(noAtoms):

    if isinstance(X_5[i], (int, float)):
        newX_5.append(X_5[i] - X_5[0])

    if isinstance(Y_5[i], (int, float)):
        newY_5.append(Y_5[i] - Y_5[0])

    if isinstance(Z_5[i], (int, float)):
        newZ_5.append(Z_5[i] - Z_5[0])
#(atom-center of mass) distances
myDistances_5 = np.sqrt([(newX_5[i]**2) + (newY_5[i]**2)+ (newZ_5[i]**2) for i in range(noAtoms)])
mySquaredDistances_5 = [] 
for x in myDistances_5:
    mySquaredDistances_5.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_5 = [mySquaredDistances_5[i] * myMasses[i] for i in range(noAtoms)]
mySquare_5Sum = sum(myMassWeightedSqDists_5)
# Radius of gyration of the structure
radiusOfGyration_5 = np.sqrt(mySquare_5Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_5 = ', sum(myDistances_5)/(noAtoms), 'nm')   
print('Radius of Gyration_5 is:', radiusOfGyration_5, 'nm')
#(atom-center of mass) distances
print('distance_5 = ', myDistances_5)


positions_6 = myPdb.getPositions(frame=5) #get the positions for the frame in nanometer
X_6 = [float(coord[0]/nanometer) for coord in positions_6]
Y_6 = [float(coord[1]/nanometer) for coord in positions_6]
Z_6 = [float(coord[2]/nanometer) for coord in positions_6]
newX_6 = [] #The coordinates of the atoms with respect to center of mass
newY_6 = []
newZ_6 = []
for i in range(noAtoms):

    if isinstance(X_6[i], (int, float)):
        newX_6.append(X_6[i] - X_6[0])

    if isinstance(Y_6[i], (int, float)):
        newY_6.append(Y_6[i] - Y_6[0])

    if isinstance(Z_6[i], (int, float)):
        newZ_6.append(Z_6[i] - Z_6[0])
#(atom-center of mass) distances
myDistances_6 = np.sqrt([(newX_6[i]**2) + (newY_6[i]**2)+ (newZ_6[i]**2) for i in range(noAtoms)])
mySquaredDistances_6 = [] 
for x in myDistances_6:
    mySquaredDistances_6.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_6 = [mySquaredDistances_6[i] * myMasses[i] for i in range(noAtoms)]
mySquare_6Sum = sum(myMassWeightedSqDists_6)
# Radius of gyration of the structure
radiusOfGyration_6 = np.sqrt(mySquare_6Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_6 = ', sum(myDistances_6)/(noAtoms), 'nm')   
print('Radius of Gyration_6 is:', radiusOfGyration_6, 'nm')
#(atom-center of mass) distances
print('distance_6 = ', myDistances_6)


positions_7 = myPdb.getPositions(frame=6) #get the positions for the frame in nanometer
X_7 = [float(coord[0]/nanometer) for coord in positions_7]
Y_7 = [float(coord[1]/nanometer) for coord in positions_7]
Z_7 = [float(coord[2]/nanometer) for coord in positions_7]
newX_7 = [] #The coordinates of the atoms with respect to center of mass
newY_7 = []
newZ_7 = []
for i in range(noAtoms):

    if isinstance(X_7[i], (int, float)):
        newX_7.append(X_7[i] - X_7[0])

    if isinstance(Y_7[i], (int, float)):
        newY_7.append(Y_7[i] - Y_7[0])

    if isinstance(Z_7[i], (int, float)):
        newZ_7.append(Z_7[i] - Z_7[0])
#(atom-center of mass) distances
myDistances_7 = np.sqrt([(newX_7[i]**2) + (newY_7[i]**2)+ (newZ_7[i]**2) for i in range(noAtoms)])
mySquaredDistances_7 = [] 
for x in myDistances_7:
    mySquaredDistances_7.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_7 = [mySquaredDistances_7[i] * myMasses[i] for i in range(noAtoms)]
mySquare_7Sum = sum(myMassWeightedSqDists_7)
# Radius of gyration of the structure
radiusOfGyration_7 = np.sqrt(mySquare_7Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_7 = ', sum(myDistances_7)/(noAtoms), 'nm')   
print('Radius of Gyration_7 is:', radiusOfGyration_7, 'nm')
#(atom-center of mass) distances
print('distance_7 = ', myDistances_7)


positions_8 = myPdb.getPositions(frame=7) #get the positions for the frame in nanometer
X_8 = [float(coord[0]/nanometer) for coord in positions_8]
Y_8 = [float(coord[1]/nanometer) for coord in positions_8]
Z_8 = [float(coord[2]/nanometer) for coord in positions_8]
newX_8 = [] #The coordinates of the atoms with respect to center of mass
newY_8 = []
newZ_8 = []
for i in range(noAtoms):

    if isinstance(X_8[i], (int, float)):
        newX_8.append(X_8[i] - X_8[0])

    if isinstance(Y_8[i], (int, float)):
        newY_8.append(Y_8[i] - Y_8[0])

    if isinstance(Z_8[i], (int, float)):
        newZ_8.append(Z_8[i] - Z_8[0])
#(atom-center of mass) distances
myDistances_8 = np.sqrt([(newX_8[i]**2) + (newY_8[i]**2)+ (newZ_8[i]**2) for i in range(noAtoms)])
mySquaredDistances_8 = [] 
for x in myDistances_8:
    mySquaredDistances_8.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_8 = [mySquaredDistances_8[i] * myMasses[i] for i in range(noAtoms)]
mySquare_8Sum = sum(myMassWeightedSqDists_8)
# Radius of gyration of the structure
radiusOfGyration_8 = np.sqrt(mySquare_8Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_8 = ', sum(myDistances_8)/(noAtoms), 'nm')   
print('Radius of Gyration_8 is:', radiusOfGyration_8, 'nm')
#(atom-center of mass) distances
print('distance_8 = ', myDistances_8)


positions_9 = myPdb.getPositions(frame=8) #get the positions for the frame in nanometer
X_9 = [float(coord[0]/nanometer) for coord in positions_9]
Y_9 = [float(coord[1]/nanometer) for coord in positions_9]
Z_9 = [float(coord[2]/nanometer) for coord in positions_9]
newX_9 = [] #The coordinates of the atoms with respect to center of mass
newY_9 = []
newZ_9 = []
for i in range(noAtoms):

    if isinstance(X_9[i], (int, float)):
        newX_9.append(X_9[i] - X_9[0])

    if isinstance(Y_9[i], (int, float)):
        newY_9.append(Y_9[i] - Y_9[0])

    if isinstance(Z_9[i], (int, float)):
        newZ_9.append(Z_9[i] - Z_9[0])
#(atom-center of mass) distances
myDistances_9 = np.sqrt([(newX_9[i]**2) + (newY_9[i]**2)+ (newZ_9[i]**2) for i in range(noAtoms)])
mySquaredDistances_9 = [] 
for x in myDistances_9:
    mySquaredDistances_9.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_9 = [mySquaredDistances_9[i] * myMasses[i] for i in range(noAtoms)]
mySquare_9Sum = sum(myMassWeightedSqDists_9)
# Radius of gyration of the structure
radiusOfGyration_9 = np.sqrt(mySquare_9Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_9 = ', sum(myDistances_9)/(noAtoms), 'nm')   
print('Radius of Gyration_9 is:', radiusOfGyration_9, 'nm')
#(atom-center of mass) distances
print('distance_9 = ', myDistances_9)


positions_10 = myPdb.getPositions(frame=9) #get the positions for the frame in nanometer
X_10 = [float(coord[0]/nanometer) for coord in positions_10]
Y_10 = [float(coord[1]/nanometer) for coord in positions_10]
Z_10 = [float(coord[2]/nanometer) for coord in positions_10]
newX_10 = [] #The coordinates of the atoms with respect to center of mass
newY_10 = []
newZ_10 = []
for i in range(noAtoms):

    if isinstance(X_10[i], (int, float)):
        newX_10.append(X_10[i] - X_10[0])

    if isinstance(Y_10[i], (int, float)):
        newY_10.append(Y_10[i] - Y_10[0])

    if isinstance(Z_10[i], (int, float)):
        newZ_10.append(Z_10[i] - Z_10[0])
#(atom-center of mass) distances
myDistances_10 = np.sqrt([(newX_10[i]**2) + (newY_10[i]**2)+ (newZ_10[i]**2) for i in range(noAtoms)])
mySquaredDistances_10 = [] 
for x in myDistances_10:
    mySquaredDistances_10.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_10 = [mySquaredDistances_10[i] * myMasses[i] for i in range(noAtoms)]
mySquare_10Sum = sum(myMassWeightedSqDists_10)
# Radius of gyration of the structure
radiusOfGyration_10 = np.sqrt(mySquare_10Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_10 = ', sum(myDistances_10)/(noAtoms), 'nm')   
print('Radius of Gyration_10 is:', radiusOfGyration_10, 'nm')
#(atom-center of mass) distances
print('distance_10 = ', myDistances_10)


positions_11 = myPdb.getPositions(frame=10) #get the positions for the frame in nanometer
X_11 = [float(coord[0]/nanometer) for coord in positions_11]
Y_11 = [float(coord[1]/nanometer) for coord in positions_11]
Z_11 = [float(coord[2]/nanometer) for coord in positions_11]
newX_11 = [] #The coordinates of the atoms with respect to center of mass
newY_11 = []
newZ_11 = []
for i in range(noAtoms):

    if isinstance(X_11[i], (int, float)):
        newX_11.append(X_11[i] - X_11[0])

    if isinstance(Y_11[i], (int, float)):
        newY_11.append(Y_11[i] - Y_11[0])

    if isinstance(Z_11[i], (int, float)):
        newZ_11.append(Z_11[i] - Z_11[0])
#(atom-center of mass) distances
myDistances_11 = np.sqrt([(newX_11[i]**2) + (newY_11[i]**2)+ (newZ_11[i]**2) for i in range(noAtoms)])
mySquaredDistances_11 = [] 
for x in myDistances_11:
    mySquaredDistances_11.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_11 = [mySquaredDistances_11[i] * myMasses[i] for i in range(noAtoms)]
mySquare_11Sum = sum(myMassWeightedSqDists_11)
# Radius of gyration of the structure
radiusOfGyration_11 = np.sqrt(mySquare_11Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_11 = ', sum(myDistances_11)/(noAtoms), 'nm')   
print('Radius of Gyration_11 is:', radiusOfGyration_11, 'nm')
#(atom-center of mass) distances
print('distance_11 = ', myDistances_11)

positions_12 = myPdb.getPositions(frame=11) #get the positions for the frame in nanometer
X_12 = [float(coord[0]/nanometer) for coord in positions_12]
Y_12 = [float(coord[1]/nanometer) for coord in positions_12]
Z_12 = [float(coord[2]/nanometer) for coord in positions_12]
newX_12 = [] #The coordinates of the atoms with respect to center of mass
newY_12 = []
newZ_12 = []
for i in range(noAtoms):

    if isinstance(X_12[i], (int, float)):
        newX_12.append(X_12[i] - X_12[0])

    if isinstance(Y_12[i], (int, float)):
        newY_12.append(Y_12[i] - Y_12[0])

    if isinstance(Z_12[i], (int, float)):
        newZ_12.append(Z_12[i] - Z_12[0])
#(atom-center of mass) distances
myDistances_12 = np.sqrt([(newX_12[i]**2) + (newY_12[i]**2)+ (newZ_12[i]**2) for i in range(noAtoms)])
mySquaredDistances_12 = [] 
for x in myDistances_12:
    mySquaredDistances_12.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_12 = [mySquaredDistances_12[i] * myMasses[i] for i in range(noAtoms)]
mySquare_12Sum = sum(myMassWeightedSqDists_12)
# Radius of gyration of the structure
radiusOfGyration_12 = np.sqrt(mySquare_12Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_12 = ', sum(myDistances_12)/(noAtoms), 'nm')   
print('Radius of Gyration_12 is:', radiusOfGyration_12, 'nm')
#(atom-center of mass) distances
print('distance_12 = ', myDistances_12)


positions_13 = myPdb.getPositions(frame=12) #get the positions for the frame in nanometer
X_13 = [float(coord[0]/nanometer) for coord in positions_13]
Y_13 = [float(coord[1]/nanometer) for coord in positions_13]
Z_13 = [float(coord[2]/nanometer) for coord in positions_13]
newX_13 = [] #The coordinates of the atoms with respect to center of mass
newY_13 = []
newZ_13 = []
for i in range(noAtoms):

    if isinstance(X_13[i], (int, float)):
        newX_13.append(X_13[i] - X_13[0])

    if isinstance(Y_13[i], (int, float)):
        newY_13.append(Y_13[i] - Y_13[0])

    if isinstance(Z_13[i], (int, float)):
        newZ_13.append(Z_13[i] - Z_13[0])
#(atom-center of mass) distances
myDistances_13 = np.sqrt([(newX_13[i]**2) + (newY_13[i]**2)+ (newZ_13[i]**2) for i in range(noAtoms)])
mySquaredDistances_13 = [] 
for x in myDistances_13:
    mySquaredDistances_13.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_13 = [mySquaredDistances_13[i] * myMasses[i] for i in range(noAtoms)]
mySquare_13Sum = sum(myMassWeightedSqDists_13)
# Radius of gyration of the structure
radiusOfGyration_13 = np.sqrt(mySquare_13Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_13 = ', sum(myDistances_13)/(noAtoms), 'nm')   
print('Radius of Gyration_13 is:', radiusOfGyration_13, 'nm')
#(atom-center of mass) distances
print('distance_13 = ', myDistances_13)


positions_14 = myPdb.getPositions(frame=13) #get the positions for the frame in nanometer
X_14 = [float(coord[0]/nanometer) for coord in positions_14]
Y_14 = [float(coord[1]/nanometer) for coord in positions_14]
Z_14 = [float(coord[2]/nanometer) for coord in positions_14]
newX_14 = [] #The coordinates of the atoms with respect to center of mass
newY_14 = []
newZ_14 = []
for i in range(noAtoms):
    if isinstance(X_14[i], (int, float)):
        newX_14.append(X_14[i] - X_14[0])

    if isinstance(Y_14[i], (int, float)):
        newY_14.append(Y_14[i] - Y_14[0])

    if isinstance(Z_14[i], (int, float)):
        newZ_14.append(Z_14[i] - Z_14[0])

#(atom-center of mass) distances
myDistances_14 = np.sqrt([(newX_14[i]**2) + (newY_14[i]**2)+ (newZ_14[i]**2) for i in range(noAtoms)])
mySquaredDistances_14 = [] 
for x in myDistances_14:
    mySquaredDistances_14.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_14 = [mySquaredDistances_14[i] * myMasses[i] for i in range(noAtoms)]
mySquare_14Sum = sum(myMassWeightedSqDists_14)
# Radius of gyration of the structure
radiusOfGyration_14 = np.sqrt(mySquare_14Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_14 = ', sum(myDistances_14)/(noAtoms), 'nm')   
print('Radius of Gyration_14 is:', radiusOfGyration_14, 'nm')
#(atom-center of mass) distances
print('distance_14 = ', myDistances_14)


positions_15 = myPdb.getPositions(frame=14) #get the positions for the frame in nanometer
X_15 = [float(coord[0]/nanometer) for coord in positions_15]
Y_15 = [float(coord[1]/nanometer) for coord in positions_15]
Z_15 = [float(coord[2]/nanometer) for coord in positions_15]
newX_15 = [] #The coordinates of the atoms with respect to center of mass
newY_15 = []
newZ_15 = []
for i in range(noAtoms):

    if isinstance(X_15[i], (int, float)):
        newX_15.append(X_15[i] - X_15[0])

    if isinstance(Y_15[i], (int, float)):
        newY_15.append(Y_15[i] - Y_15[0])

    if isinstance(Z_15[i], (int, float)):
        newZ_15.append(Z_15[i] - Z_15[0])
#(atom-center of mass) distances
myDistances_15 = np.sqrt([(newX_15[i]**2) + (newY_15[i]**2)+ (newZ_15[i]**2) for i in range(noAtoms)])
mySquaredDistances_15 = [] 
for x in myDistances_15:
    mySquaredDistances_15.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_15 = [mySquaredDistances_15[i] * myMasses[i] for i in range(noAtoms)]
mySquare_15Sum = sum(myMassWeightedSqDists_15)
# Radius of gyration of the structure
radiusOfGyration_15 = np.sqrt(mySquare_15Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_15 = ', sum(myDistances_15)/(noAtoms), 'nm')   
print('Radius of Gyration_15 is:', radiusOfGyration_15, 'nm')
#(atom-center of mass) distances
print('distance_15 = ', myDistances_15)


positions_16 = myPdb.getPositions(frame=15) #get the positions for the frame in nanometer
X_16 = [float(coord[0]/nanometer) for coord in positions_16]
Y_16 = [float(coord[1]/nanometer) for coord in positions_16]
Z_16 = [float(coord[2]/nanometer) for coord in positions_16]
newX_16 = [] #The coordinates of the atoms with respect to center of mass
newY_16 = []
newZ_16 = []
for i in range(noAtoms):

    if isinstance(X_16[i], (int, float)):
        newX_16.append(X_16[i] - X_16[0])

    if isinstance(Y_16[i], (int, float)):
        newY_16.append(Y_16[i] - Y_16[0])

    if isinstance(Z_16[i], (int, float)):
        newZ_16.append(Z_16[i] - Z_16[0])
#(atom-center of mass) distances
myDistances_16 = np.sqrt([(newX_16[i]**2) + (newY_16[i]**2)+ (newZ_16[i]**2) for i in range(noAtoms)])
mySquaredDistances_16 = [] 
for x in myDistances_16:
    mySquaredDistances_16.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_16 = [mySquaredDistances_16[i] * myMasses[i] for i in range(noAtoms)]
mySquare_16Sum = sum(myMassWeightedSqDists_16)
# Radius of gyration of the structure
radiusOfGyration_16 = np.sqrt(mySquare_16Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_16 = ', sum(myDistances_16)/(noAtoms), 'nm')   
print('Radius of Gyration_16 is:', radiusOfGyration_16, 'nm')
#(atom-center of mass) distances
print('distance_16 = ', myDistances_16)


positions_17 = myPdb.getPositions(frame=16) #get the positions for the frame in nanometer
X_17 = [float(coord[0]/nanometer) for coord in positions_17]
Y_17 = [float(coord[1]/nanometer) for coord in positions_17]
Z_17 = [float(coord[2]/nanometer) for coord in positions_17]
newX_17 = [] #The coordinates of the atoms with respect to center of mass
newY_17 = []
newZ_17 = []
for i in range(noAtoms):

    if isinstance(X_17[i], (int, float)):
        newX_17.append(X_17[i] - X_17[0])

    if isinstance(Y_17[i], (int, float)):
        newY_17.append(Y_17[i] - Y_17[0])

    if isinstance(Z_17[i], (int, float)):
        newZ_17.append(Z_17[i] - Z_17[0])
#(atom-center of mass) distances
myDistances_17 = np.sqrt([(newX_17[i]**2) + (newY_17[i]**2)+ (newZ_17[i]**2) for i in range(noAtoms)])
mySquaredDistances_17 = [] 
for x in myDistances_17:
    mySquaredDistances_17.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_17 = [mySquaredDistances_17[i] * myMasses[i] for i in range(noAtoms)]
mySquare_17Sum = sum(myMassWeightedSqDists_17)
# Radius of gyration of the structure
radiusOfGyration_17 = np.sqrt(mySquare_17Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_17 = ', sum(myDistances_17)/(noAtoms), 'nm')   
print('Radius of Gyration_17 is:', radiusOfGyration_17, 'nm')
#(atom-center of mass) distances
print('distance_17 = ', myDistances_17)


positions_18 = myPdb.getPositions(frame=17) #get the positions for the frame in nanometer
X_18 = [float(coord[0]/nanometer) for coord in positions_18]
Y_18 = [float(coord[1]/nanometer) for coord in positions_18]
Z_18 = [float(coord[2]/nanometer) for coord in positions_18]
newX_18 = [] #The coordinates of the atoms with respect to center of mass
newY_18 = []
newZ_18 = []
for i in range(noAtoms):

    if isinstance(X_18[i], (int, float)):
        newX_18.append(X_18[i] - X_18[0])

    if isinstance(Y_18[i], (int, float)):
        newY_18.append(Y_18[i] - Y_18[0])

    if isinstance(Z_18[i], (int, float)):
        newZ_18.append(Z_18[i] - Z_18[0])
#(atom-center of mass) distances
myDistances_18 = np.sqrt([(newX_18[i]**2) + (newY_18[i]**2)+ (newZ_18[i]**2) for i in range(noAtoms)])
mySquaredDistances_18 = [] 
for x in myDistances_18:
    mySquaredDistances_18.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_18 = [mySquaredDistances_18[i] * myMasses[i] for i in range(noAtoms)]
mySquare_18Sum = sum(myMassWeightedSqDists_18)
# Radius of gyration of the structure
radiusOfGyration_18 = np.sqrt(mySquare_18Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_18 = ', sum(myDistances_18)/(noAtoms), 'nm')   
print('Radius of Gyration_18 is:', radiusOfGyration_18, 'nm')
#(atom-center of mass) distances
print('distance_18 = ', myDistances_18)


positions_19 = myPdb.getPositions(frame=18) #get the positions for the frame in nanometer
X_19 = [float(coord[0]/nanometer) for coord in positions_19]
Y_19 = [float(coord[1]/nanometer) for coord in positions_19]
Z_19 = [float(coord[2]/nanometer) for coord in positions_19]
newX_19 = [] #The coordinates of the atoms with respect to center of mass
newY_19 = []
newZ_19 = []
for i in range(noAtoms):

    if isinstance(X_19[i], (int, float)):
        newX_19.append(X_19[i] - X_19[0])

    if isinstance(Y_19[i], (int, float)):
        newY_19.append(Y_19[i] - Y_19[0])

    if isinstance(Z_19[i], (int, float)):
        newZ_19.append(Z_19[i] - Z_19[0])
#(atom-center of mass) distances
myDistances_19 = np.sqrt([(newX_19[i]**2) + (newY_19[i]**2)+ (newZ_19[i]**2) for i in range(noAtoms)])
mySquaredDistances_19 = [] 
for x in myDistances_19:
    mySquaredDistances_19.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_19 = [mySquaredDistances_19[i] * myMasses[i] for i in range(noAtoms)]
mySquare_19Sum = sum(myMassWeightedSqDists_19)
# Radius of gyration of the structure
radiusOfGyration_19 = np.sqrt(mySquare_19Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_19 = ', sum(myDistances_19)/(noAtoms), 'nm')   
print('Radius of Gyration_19 is:', radiusOfGyration_19, 'nm')
#(atom-center of mass) distances
print('distance_19 = ', myDistances_19)

positions_20 = myPdb.getPositions(frame=19) #get the positions for the frame in nanometer
X_20 = [float(coord[0]/nanometer) for coord in positions_20]
Y_20 = [float(coord[1]/nanometer) for coord in positions_20]
Z_20 = [float(coord[2]/nanometer) for coord in positions_20]
newX_20 = [] #The coordinates of the atoms with respect to center of mass
newY_20 = []
newZ_20 = []
for i in range(noAtoms):

    if isinstance(X_20[i], (int, float)):
        newX_20.append(X_20[i] - X_20[0])

    if isinstance(Y_20[i], (int, float)):
        newY_20.append(Y_20[i] - Y_20[0])

    if isinstance(Z_20[i], (int, float)):
        newZ_20.append(Z_20[i] - Z_20[0])
#(atom-center of mass) distances
myDistances_20 = np.sqrt([(newX_20[i]**2) + (newY_20[i]**2)+ (newZ_20[i]**2) for i in range(noAtoms)])
mySquaredDistances_20 = [] 
for x in myDistances_20:
    mySquaredDistances_20.append(x**2)
#The square of (atom-center of mass) * atom mass
myMassWeightedSqDists_20 = [mySquaredDistances_20[i] * myMasses[i] for i in range(noAtoms)]
mySquare_20Sum = sum(myMassWeightedSqDists_20)
# Radius of gyration of the structure
radiusOfGyration_20 = np.sqrt(mySquare_20Sum/totalMass)
#The average of(atom-center of mass) distance of the structure
print('Average distance_20 = ', sum(myDistances_20)/(noAtoms), 'nm')   
print('Radius of Gyration_20 is:', radiusOfGyration_20, 'nm')
#(atom-center of mass) distances
print('distance_20 = ', myDistances_20)


averageDistances_30M = [sum(myDistances_1)/noAtoms, sum(myDistances_2)/noAtoms, sum(myDistances_3)/noAtoms, sum(myDistances_4)/noAtoms, sum(myDistances_5)/noAtoms, sum(myDistances_6)/noAtoms, sum(myDistances_7)/noAtoms, sum(myDistances_8)/noAtoms, sum(myDistances_9)/noAtoms, sum(myDistances_10)/noAtoms, sum(myDistances_11)/noAtoms, sum(myDistances_12)/noAtoms, sum(myDistances_13)/noAtoms, sum(myDistances_14)/noAtoms, sum(myDistances_15)/noAtoms, sum(myDistances_16)/noAtoms, sum(myDistances_17)/noAtoms, sum(myDistances_18)/noAtoms, sum(myDistances_19)/noAtoms, sum(myDistances_20)/noAtoms]

radiiOfGyration_30M = [radiusOfGyration_1, radiusOfGyration_2, radiusOfGyration_3, radiusOfGyration_4, radiusOfGyration_5, radiusOfGyration_6, radiusOfGyration_7, radiusOfGyration_8, radiusOfGyration_9, radiusOfGyration_10, radiusOfGyration_11, radiusOfGyration_12, radiusOfGyration_20, radiusOfGyration_14, radiusOfGyration_15, radiusOfGyration_16, radiusOfGyration_17, radiusOfGyration_18, radiusOfGyration_19, radiusOfGyration_20]

 # Histogram plots for the last ten structures to show how atoms are distributed through the '10 shells' by structures

x_shells = ['first','second','third','fourth','fifth','sixth','seventh','eighth','ninth','tenth']

plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_11, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 11th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_11, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_11 = count[0]
A_11 = count[1]
W_11 = [A_11[1]-A_11[0], A_11[2]-A_11[1], A_11[3]-A_11[2], A_11[4]-A_11[3], A_11[5]-A_11[4], A_11[6]-A_11[5], A_11[7]-A_11[6], A_11[8]-A_11[7], A_11[9]-A_11[8], A_11[10]-A_11[9]]
L_11 = B_11.tolist()

print("Width",W_11, "Length", L_11)
print(len(W_11),len(L_11))
volumes_11 = []
for i in range(len(L_11)):
    v_11 = W_11[i]*L_11[i]
    volumes_11.append(v_11)

Densities_11=[]

for i in range(len(L_11)):
    d_11 = L_11[i]/volumes_11[i]
    Densities_11.append(d_11)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_11, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_11)
print('Volumes of every shell_11th strcture@30M run:', volumes_11)
print('The Densities of the shells:', Densities_11)



plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_12, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 12th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_12, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_12 = count[0]
A_12 = count[1]
W_12 = [A_12[1]-A_12[0], A_12[2]-A_12[1], A_12[3]-A_12[2], A_12[4]-A_12[3], A_12[5]-A_12[4], A_12[6]-A_12[5], A_12[7]-A_12[6], A_12[8]-A_12[7], A_12[9]-A_12[8], A_12[10]-A_12[9]]
L_12 = B_12.tolist()

print("Width",W_12, "Length", L_12)
print(len(W_12),len(L_12))
volumes_12 = []
for i in range(len(L_12)):
    v_12 = W_12[i]*L_12[i]
    volumes_12.append(v_12)

Densities_12=[]

for i in range(len(L_12)):
    d_12 = L_12[i]/volumes_12[i]
    Densities_12.append(d_12)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_12, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_12)
print('Volumes of every shell_12th strcture@30M run:', volumes_12)
print('The Densities of the shells:', Densities_12)


plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_13, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 13th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_13, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_13 = count[0]
A_13 = count[1]
W_13 = [A_13[1]-A_13[0], A_13[2]-A_13[1], A_13[3]-A_13[2], A_13[4]-A_13[3], A_13[5]-A_13[4], A_13[6]-A_13[5], A_13[7]-A_13[6], A_13[8]-A_13[7], A_13[9]-A_13[8], A_13[10]-A_13[9]]
L_13 = B_13.tolist()

print("Width",W_13, "Length", L_13)
print(len(W_13),len(L_13))
volumes_13 = []
for i in range(len(L_13)):
    v_13 = W_13[i]*L_13[i]
    volumes_13.append(v_13)

Densities_13=[]

for i in range(len(L_13)):
    d_13 = L_13[i]/volumes_13[i]
    Densities_13.append(d_13)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_13, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_13)
print('Volumes of every shell_13th strcture@30M run:', volumes_13)
print('The Densities of the shells:', Densities_13)


plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_14, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 14th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_14, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_14 = count[0]
A_14 = count[1]
W_14 = [A_14[1]-A_14[0], A_14[2]-A_14[1], A_14[3]-A_14[2], A_14[4]-A_14[3], A_14[5]-A_14[4], A_14[6]-A_14[5], A_14[7]-A_14[6], A_14[8]-A_14[7], A_14[9]-A_14[8], A_14[10]-A_14[9]]
L_14 = B_14.tolist()

print("Width",W_14, "Length", L_14)
print(len(W_14),len(L_14))
volumes_14 = []
for i in range(len(L_14)):
    v_14 = W_14[i]*L_14[i]
    volumes_14.append(v_14)

Densities_14=[]

for i in range(len(L_14)):
    d_14 = L_14[i]/volumes_14[i]
    Densities_14.append(d_14)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_14, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_14)
print('Volumes of every shell_14th strcture@30M run:', volumes_14)
print('The Densities of the shells:', Densities_14)



plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_15, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 15th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_15, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_15 = count[0]
A_15 = count[1]
W_15 = [A_15[1]-A_15[0], A_15[2]-A_15[1], A_15[3]-A_15[2], A_15[4]-A_15[3], A_15[5]-A_15[4], A_15[6]-A_15[5], A_15[7]-A_15[6], A_15[8]-A_15[7], A_15[9]-A_15[8], A_15[10]-A_15[9]]
L_15 = B_15.tolist()

print("Width",W_15, "Length", L_15)
print(len(W_15),len(L_15))
volumes_15 = []
for i in range(len(L_15)):
    v_15 = W_15[i]*L_15[i]
    volumes_15.append(v_15)

Densities_15=[]

for i in range(len(L_15)):
    d_15 = L_15[i]/volumes_15[i]
    Densities_15.append(d_15)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_15, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_15)
print('Volumes of every shell_15th strcture@30M run:', volumes_15)
print('The Densities of the shells:', Densities_15)


plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_16, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 16th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_16, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_16 = count[0]
A_16 = count[1]
W_16 = [A_16[1]-A_16[0], A_16[2]-A_16[1], A_16[3]-A_16[2], A_16[4]-A_16[3], A_16[5]-A_16[4], A_16[6]-A_16[5], A_16[7]-A_16[6], A_16[8]-A_16[7], A_16[9]-A_16[8], A_16[10]-A_16[9]]
L_16 = B_16.tolist()

print("Width",W_16, "Length", L_16)
print(len(W_16),len(L_16))
volumes_16 = []
for i in range(len(L_16)):
    v_16 = W_16[i]*L_16[i]
    volumes_16.append(v_16)

Densities_16=[]

for i in range(len(L_16)):
    d_16 = L_16[i]/volumes_16[i]
    Densities_16.append(d_16)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_16, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_16)
print('Volumes of every shell_16th strcture@30M run:', volumes_16)
print('The Densities of the shells:', Densities_16)


plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_17, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 17th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_17, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_17 = count[0]
A_17 = count[1]
W_17 = [A_17[1]-A_17[0], A_17[2]-A_17[1], A_17[3]-A_17[2], A_17[4]-A_17[3], A_17[5]-A_17[4], A_17[6]-A_17[5], A_17[7]-A_17[6], A_17[8]-A_17[7], A_17[9]-A_17[8], A_17[10]-A_17[9]]
L_17 = B_17.tolist()

print("Width",W_17, "Length", L_17)
print(len(W_17),len(L_17))
volumes_17 = []
for i in range(len(L_17)):
    v_17 = W_17[i]*L_17[i]
    volumes_17.append(v_17)

Densities_17=[]

for i in range(len(L_17)):
    d_17 = L_17[i]/volumes_17[i]
    Densities_17.append(d_17)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_17, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_17)
print('Volumes of every shell_17th strcture@30M run:', volumes_17)
print('The Densities of the shells:', Densities_17)


plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_18, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 18th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_18, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_18 = count[0]
A_18 = count[1]
W_18 = [A_18[1]-A_18[0], A_18[2]-A_18[1], A_18[3]-A_18[2], A_18[4]-A_18[3], A_18[5]-A_18[4], A_18[6]-A_18[5], A_18[7]-A_18[6], A_18[8]-A_18[7], A_18[9]-A_18[8], A_18[10]-A_18[9]]
L_18 = B_18.tolist()

print("Width",W_18, "Length", L_18)
print(len(W_18),len(L_18))
volumes_18 = []
for i in range(len(L_18)):
    v_18 = W_18[i]*L_18[i]
    volumes_18.append(v_18)

Densities_18=[]

for i in range(len(L_18)):
    d_18 = L_18[i]/volumes_18[i]
    Densities_18.append(d_18)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_18, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_18)
print('Volumes of every shell_18th strcture@30M run:', volumes_18)
print('The Densities of the shells:', Densities_18)



plt.figure(figsize=(10,8))
_values, bin_edges, _rectangles = plt.hist(myDistances_19, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 19th structure @ 30M steps simulation as ten shells', fontsize=20, pad=5)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_19, bins=10)
plt.margins(x=0.012)
plt.tight_layout()
plt.show()

B_19 = count[0]
A_19 = count[1]
W_19 = [A_19[1]-A_19[0], A_19[2]-A_19[1], A_19[3]-A_19[2], A_19[4]-A_19[3], A_19[5]-A_19[4], A_19[6]-A_19[5], A_19[7]-A_19[6], A_19[8]-A_19[7], A_19[9]-A_19[8], A_19[10]-A_19[9]]
L_19 = B_19.tolist()

print("Width",W_19, "Length", L_19)
print(len(W_19),len(L_19))
volumes_19 = []
for i in range(len(L_19)):
    v_19 = W_19[i]*L_19[i]
    volumes_19.append(v_19)

Densities_19=[]

for i in range(len(L_19)):
    d_19 = L_19[i]/volumes_19[i]
    Densities_19.append(d_19)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_19, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_19)
print('Volumes of every shell_19th strcture@30M run:', volumes_19)
print('The Densities of the shells:', Densities_19)


plt.figure(figsize=(15,20))
_values, bin_edges, _rectangles = plt.hist(myDistances_20, bins=10, edgecolor='k', facecolor='b')
plt.xticks((bin_edges[:-1] + bin_edges[1:]) / 2,
           ['first', 'second', 'third', 'fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth'])
plt.title('The atoms distribution in the 20th structure @ 30M steps simulation as ten shells', pad=6, weight="bold", fontsize=20)
plt.ylabel('Number of atoms', fontsize=15)
plt.xlabel('Shell number', fontsize=15)
count = np.histogram(myDistances_20, bins=10)
plt.margins(x=0.012)
plt.tight_layout(pad=15)
plt.show()

B_20 = count[0]
A_20 = count[1]
W_20 = [A_20[1]-A_20[0], A_20[2]-A_20[1], A_20[3]-A_20[2], A_20[4]-A_20[3], A_20[5]-A_20[4], A_20[6]-A_20[5], A_20[7]-A_20[6], A_20[8]-A_20[7], A_20[9]-A_20[8], A_20[10]-A_20[9]]
L_20 = B_20.tolist()

print("Width",W_20, "Length", L_20)
print(len(W_20),len(L_20))
volumes_20 = []
for i in range(len(L_20)):
    v_20 = W_20[i]*L_20[i]
    volumes_20.append(v_20)

Densities_20=[]

for i in range(len(L_20)):
    d_20 = L_20[i]/volumes_20[i]
    Densities_20.append(d_20)
print('The histogram:- ', 'shows ten shells per structure', " Widths of th shells are:", W_20, "\n", "Heights of the shells, also it shows the number of atoms in the shell", L_20)
print('Volumes of every shell_20th strcture@30M run:', volumes_20)
print('The Densities of the shells:', Densities_20)

Mean_noAtoms = [statistics.mean(L_11),statistics.mean(L_12),statistics.mean(L_13),statistics.mean(L_14),statistics.mean(L_15),statistics.mean(L_16),statistics.mean(L_17),statistics.mean(L_18),statistics.mean(L_19),statistics.mean(L_20)]
print('Mean_noAtoms',Mean_noAtoms)
#
SD_noAtoms = [statistics.stdev(L_11),statistics.stdev(L_12),statistics.stdev(L_13),statistics.stdev(L_14),statistics.stdev(L_15),statistics.stdev(L_16),statistics.stdev(L_17),statistics.stdev(L_18),statistics.stdev(L_19),statistics.stdev(L_20)]
print('SD_noAtoms',SD_noAtoms)


Mean_densities = [statistics.mean(Densities_11),statistics.mean(Densities_12),statistics.mean(Densities_13),statistics.mean(Densities_14),statistics.mean(Densities_15),statistics.mean(Densities_16),statistics.mean(Densities_17),statistics.mean(Densities_18),statistics.mean(Densities_19),statistics.mean(Densities_20)]
print('Mean_densities',Mean_densities)

SD_densities = [statistics.stdev(Densities_11),statistics.stdev(Densities_12),statistics.stdev(Densities_13),statistics.stdev(Densities_14),statistics.stdev(Densities_15),statistics.stdev(Densities_16),statistics.stdev(Densities_17),statistics.stdev(Densities_18),statistics.stdev(Densities_19),statistics.stdev(Densities_20)]
print('SD_densities',SD_densities)


Mean_volumes = [statistics.mean(volumes_11),statistics.mean(volumes_12),statistics.mean(volumes_13),statistics.mean(volumes_14),statistics.mean(volumes_15),statistics.mean(volumes_16),statistics.mean(volumes_17),statistics.mean(volumes_18),statistics.mean(volumes_19),statistics.mean(volumes_20)]
print('Mean_volumes',Mean_volumes)

SD_volumes = [statistics.stdev(volumes_11),statistics.stdev(volumes_12),statistics.stdev(volumes_13),statistics.stdev(volumes_14),statistics.stdev(volumes_15),statistics.stdev(volumes_16),statistics.stdev(volumes_17),statistics.stdev(volumes_18),statistics.stdev(volumes_19),statistics.stdev(volumes_20)]
print('SD_volumes',SD_volumes)

firstNoAtoms = ((L_11[0],L_12[0],L_13[0],L_14[0],L_15[0],L_16[0],L_17[0],L_18[0],L_19[0],L_20[0]))

secondNoAtoms = ((L_11[1],L_12[1],L_13[1],L_14[1],L_15[1],L_16[1],L_17[1],L_18[1],L_19[1],L_20[1]))

thirdNoAtoms = ((L_11[2],L_12[2],L_13[2],L_14[2],L_15[2],L_16[2],L_17[2],L_18[2],L_19[2],L_20[2]))

fourthNoAtoms = ((L_11[3],L_12[3],L_13[3],L_14[3],L_15[3],L_16[3],L_17[3],L_18[3],L_19[3],L_20[3]))

fifthNoAtoms = ((L_11[4],L_12[4],L_13[4],L_14[4],L_15[4],L_16[4],L_17[4],L_18[4],L_19[4],L_20[4]))

sixthNoAtoms = ((L_11[5],L_12[5],L_13[5],L_14[5],L_15[5],L_16[5],L_17[5],L_18[5],L_19[5],L_20[5]))

seventhNoAtoms = ((L_11[6],L_12[6],L_13[6],L_14[6],L_15[6],L_16[6],L_17[6],L_18[6],L_19[6],L_20[6]))

eighthNoAtoms = ((L_11[7],L_12[7],L_13[7],L_14[7],L_15[7],L_16[7],L_17[7],L_18[7],L_19[7],L_20[7]))

ninthNoAtoms = ((L_11[8],L_12[8],L_13[8],L_14[8],L_15[8],L_16[8],L_17[8],L_18[8],L_19[8],L_20[8]))

tenthNoAtoms = ((L_11[9],L_12[9],L_13[9],L_14[9],L_15[9],L_16[9],L_17[9],L_18[9],L_19[9],L_20[9]))

meanNoAtoms = [statistics.mean(firstNoAtoms),statistics.mean(secondNoAtoms),statistics.mean(thirdNoAtoms),statistics.mean(fourthNoAtoms),statistics.mean(fifthNoAtoms),statistics.mean(sixthNoAtoms),statistics.mean(seventhNoAtoms),statistics.mean(eighthNoAtoms),statistics.mean(ninthNoAtoms),statistics.mean(tenthNoAtoms)]
print('meanNoAtoms', meanNoAtoms)

stdevNoAtoms = [statistics.stdev(firstNoAtoms),statistics.stdev(secondNoAtoms),statistics.stdev(thirdNoAtoms),statistics.stdev(fourthNoAtoms),statistics.stdev(fifthNoAtoms),statistics.stdev(sixthNoAtoms),statistics.stdev(seventhNoAtoms),statistics.stdev(eighthNoAtoms),statistics.stdev(ninthNoAtoms),statistics.stdev(tenthNoAtoms)]
print('stdevNoAtoms', stdevNoAtoms)

firstDensities = ((Densities_11[0],Densities_12[0],Densities_13[0],Densities_14[0],Densities_15[0],Densities_16[0],Densities_17[0],Densities_18[0],Densities_19[0],Densities_20[0]))

secondDensities = ((Densities_11[1],Densities_12[1],Densities_13[1],Densities_14[1],Densities_15[1],Densities_16[1],Densities_17[1],Densities_18[1],Densities_19[1],Densities_20[1]))

thirdDensities = ((Densities_11[2],Densities_12[2],Densities_13[2],Densities_14[2],Densities_15[2],Densities_16[2],Densities_17[2],Densities_18[2],Densities_19[2],Densities_20[2]))

fourthDensities = ((Densities_11[3],Densities_12[3],Densities_13[3],Densities_14[3],Densities_15[3],Densities_16[3],Densities_17[3],Densities_18[3],Densities_19[3],Densities_20[3]))

fifthDensities = ((Densities_11[4],Densities_12[4],Densities_13[4],Densities_14[4],Densities_15[4],Densities_16[4],Densities_17[4],Densities_18[4],Densities_19[4],Densities_20[4]))

sixthDensities = ((Densities_11[5],Densities_12[5],Densities_13[5],Densities_14[5],Densities_15[5],Densities_16[5],Densities_17[5],Densities_18[5],Densities_19[5],Densities_20[5]))

seventhDensities = ((Densities_11[6],Densities_12[6],Densities_13[6],Densities_14[6],Densities_15[6],Densities_16[6],Densities_17[6],Densities_18[6],Densities_19[6],Densities_20[6]))

eighthDensities = ((Densities_11[7],Densities_12[7],Densities_13[7],Densities_14[7],Densities_15[7],Densities_16[7],Densities_17[7],Densities_18[7],Densities_19[7],Densities_20[7]))

ninthDensities = ((Densities_11[8],Densities_12[8],Densities_13[8],Densities_14[8],Densities_15[8],Densities_16[8],Densities_17[8],Densities_18[8],Densities_19[8],Densities_20[8]))

tenthDensities = ((Densities_11[9],Densities_12[9],Densities_13[9],Densities_14[9],Densities_15[9],Densities_16[9],Densities_17[9],Densities_18[9],Densities_19[9],Densities_20[9]))

meanDensities = [statistics.mean(firstDensities),statistics.mean(secondDensities),statistics.mean(thirdDensities),statistics.mean(fourthDensities),statistics.mean(fifthDensities),statistics.mean(sixthDensities),statistics.mean(seventhDensities),statistics.mean(eighthDensities),statistics.mean(ninthDensities),statistics.mean(tenthDensities)]
print('meanDensities', meanDensities)

stdevDensities = [statistics.stdev(firstDensities),statistics.stdev(secondDensities),statistics.stdev(thirdDensities),statistics.stdev(fourthDensities),statistics.stdev(fifthDensities),statistics.stdev(sixthDensities),statistics.stdev(seventhDensities),statistics.stdev(eighthDensities),statistics.stdev(ninthDensities),statistics.stdev(tenthDensities)]
print('stdevDensities', stdevDensities)


firstvolumes = ((volumes_11[0],volumes_12[0],volumes_13[0],volumes_14[0],volumes_15[0],volumes_16[0],volumes_17[0],volumes_18[0],volumes_19[0],volumes_20[0]))

secondvolumes = ((volumes_11[1],volumes_12[1],volumes_13[1],volumes_14[1],volumes_15[1],volumes_16[1],volumes_17[1],volumes_18[1],volumes_19[1],volumes_20[1]))

thirdvolumes = ((volumes_11[2],volumes_12[2],volumes_13[2],volumes_14[2],volumes_15[2],volumes_16[2],volumes_17[2],volumes_18[2],volumes_19[2],volumes_20[2]))

fourthvolumes = ((volumes_11[3],volumes_12[3],volumes_13[3],volumes_14[3],volumes_15[3],volumes_16[3],volumes_17[3],volumes_18[3],volumes_19[3],volumes_20[3]))

fifthvolumes = ((volumes_11[4],volumes_12[4],volumes_13[4],volumes_14[4],volumes_15[4],volumes_16[4],volumes_17[4],volumes_18[4],volumes_19[4],volumes_20[4]))

sixthvolumes = ((volumes_11[5],volumes_12[5],volumes_13[5],volumes_14[5],volumes_15[5],volumes_16[5],volumes_17[5],volumes_18[5],volumes_19[5],volumes_20[5]))

seventhvolumes = ((volumes_11[6],volumes_12[6],volumes_13[6],volumes_14[6],volumes_15[6],volumes_16[6],volumes_17[6],volumes_18[6],volumes_19[6],volumes_20[6]))

eighthvolumes = ((volumes_11[7],volumes_12[7],volumes_13[7],volumes_14[7],volumes_15[7],volumes_16[7],volumes_17[7],volumes_18[7],volumes_19[7],volumes_20[7]))

ninthvolumes = ((volumes_11[8],volumes_12[8],volumes_13[8],volumes_14[8],volumes_15[8],volumes_16[8],volumes_17[8],volumes_18[8],volumes_19[8],volumes_20[8]))

tenthvolumes = ((volumes_11[9],volumes_12[9],volumes_13[9],volumes_14[9],volumes_15[9],volumes_16[9],volumes_17[9],volumes_18[9],volumes_19[9],volumes_20[9]))

meanvolumes = [statistics.mean(firstvolumes),statistics.mean(secondvolumes),statistics.mean(thirdvolumes),statistics.mean(fourthvolumes),statistics.mean(fifthvolumes),statistics.mean(sixthvolumes),statistics.mean(seventhvolumes),statistics.mean(eighthvolumes),statistics.mean(ninthvolumes),statistics.mean(tenthvolumes)]
print('meanvolumes', meanvolumes)


stdevvolumes = [statistics.stdev(firstvolumes),statistics.stdev(secondvolumes),statistics.stdev(thirdvolumes),statistics.stdev(fourthvolumes),statistics.stdev(fifthvolumes),statistics.stdev(sixthvolumes),statistics.stdev(seventhvolumes),statistics.stdev(eighthvolumes),statistics.stdev(ninthvolumes),statistics.stdev(tenthvolumes)]
print('stdevvolumes', stdevvolumes)

# Build the plot
x_pos = np.arange(10)
shells = ['first','second', 'third','fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth']
fig, ax = plt.subplots(figsize=(10,10))

ax.bar(x_pos, meanDensities, yerr=stdevDensities, align='center', alpha=0.5, ecolor='black', capsize=10)
ax.set_ylabel('Atom density ($nm^{-3}$)', fontweight='bold', fontsize=15)
ax.set_xlabel('Shell number', fontweight='bold', fontsize=15)
ax.set_xticks(x_pos)
ax.set_xticklabels(shells, rotation=30, fontweight='bold', fontsize=12)
ax.set_title('Density of atoms per shell',pad=4, fontweight='bold', fontsize=20)
plt.tight_layout()
plt.show()


# Build the plot
x_pos = np.arange(10)
shells = ['first','second', 'third','fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth']
fig, ax = plt.subplots(figsize=(10,10))

ax.bar(x_pos, meanNoAtoms, yerr=stdevNoAtoms, align='center', alpha=0.5, ecolor='black', capsize=10)
ax.set_ylabel('Number of atoms ($Atoms$)', fontweight='bold', fontsize=15)
ax.set_xlabel('Shell number', fontweight='bold', fontsize=15)
ax.set_xticks(x_pos)
ax.set_xticklabels(shells, rotation=30, fontweight='bold', fontsize=12)
ax.set_title('Number of atoms per shell',pad=4, fontweight='bold', fontsize=20)
# ax.yaxis.grid(True)
plt.tight_layout()
plt.show()

# Build the plot
x_pos = np.arange(10)
shells = ['first','second', 'third','fourth', 'fifth', 'sixth', 'seventh', 'eighth', 'ninth', 'tenth']
fig, ax = plt.subplots(figsize=(10,10))

ax.bar(x_pos, meanvolumes, yerr=stdevvolumes, align='center', alpha=0.5, ecolor='black', capsize=10)
ax.set_ylabel('Volumes of individual shell ($nm^3$)', fontweight='bold', fontsize=15)
ax.set_xlabel('Shell number', fontweight='bold', fontsize=15)
ax.set_xticks(x_pos)
ax.set_xticklabels(shells, rotation=30, fontweight='bold', fontsize=12)
ax.set_title('Volumes per shell', pad=4, fontweight='bold', fontsize=20)
plt.tight_layout()
plt.show()

# Build the plot
x_pos = np.arange(10)
structures = ['11th','12th', '13th','14th', '15th', '16th', '17th', '18th', '19th', '20th']
fig, ax = plt.subplots(figsize=(10,10))

ax.bar(x_pos, Mean_noAtoms, yerr=SD_noAtoms, align='center', alpha=0.5, ecolor='black', capsize=10)
ax.set_ylabel('Number of atoms ($Atoms$)', fontweight='bold', fontsize=15)
ax.set_xlabel('Structure number', fontweight='bold', fontsize=15)
ax.set_xticks(x_pos)
ax.set_xticklabels(structures, rotation=30, fontweight='bold', fontsize=10)
ax.set_title('Number of atoms per structure', pad=4, fontweight='bold', fontsize=20)
plt.tight_layout()
plt.show()
# Build the plot
x_pos = np.arange(10)
fig, ax = plt.subplots(figsize=(10,10))

ax.bar(x_pos, Mean_densities, yerr=SD_densities, align='center', alpha=0.5, ecolor='black', capsize=10)
ax.set_ylabel('Density of atoms ($Atom/nm^3$)', fontweight='bold',fontsize=15)
ax.set_xlabel('Structure number', fontweight='bold',fontsize=15)
ax.set_xticks(x_pos)
ax.set_xticklabels(structures, rotation=30, fontweight='bold',fontsize=10)
ax.set_title('Density of atoms per structure',fontweight='bold',pad=4,fontsize=20)
# ax.yaxis.grid(True)
plt.tight_layout()
plt.show()
# Build the plot
x_pos = np.arange(10)
fig, ax = plt.subplots(figsize=(10,10))

ax.bar(x_pos, Mean_volumes, yerr=SD_volumes, align='center', alpha=0.5, ecolor='black', capsize=10)
ax.set_ylabel('Volumes of structures ($nm^3$)', fontweight='bold',fontsize=15)
ax.set_xlabel('Structure number', fontweight='bold')
ax.set_xticks(x_pos)
ax.set_xticklabels(shells, rotation=30, fontweight='bold',fontsize=15)
ax.set_title('volumes per structure', fontweight='bold',fontsize=20, pad=4)
plt.tight_layout()
plt.show()


#Writing results to docx file
doc = docx.Document()
doc.add_heading('Boltorn (20 structure @ 30M steps)',0)
intro1 = doc.add_paragraph()
intro1.add_run('The data we got from the simulation and histogram; the atom-center of mass distances for the 20 structure, radii of gyration and the average atom-center of mass per all the atoms in the structure.'+'\n').italic = True
intro1.add_run('Then we used the atom-center of mass distances for last ten structures, dividing every structure into ten shells by plotting histogram to get some details about these shells').italic = True
h1 = doc.add_paragraph('The structure number 1: ', style='Intense Quote')
d1 = doc.add_paragraph()
d1.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p1 =doc.add_paragraph('[')
for i in myDistances_1:
    p1.add_run(str(i))
    if i != myDistances_1[-1]:
        p1.add_run(',')
p1.add_run(']')
a1 = doc.add_paragraph()
a1.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a1.add_run(str(sum(myDistances_1)/(noAtoms)))

r1 = doc.add_paragraph()
r1.add_run('Radius of gyration = ').bold = True
r1.add_run(str(radiusOfGyration_1))

h2 = doc.add_paragraph('The structure number 2: ', style='Intense Quote')
d2 = doc.add_paragraph()
d2.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p2 =doc.add_paragraph('[')
for i in myDistances_2:
    p2.add_run(str(i))
    if i != myDistances_2[-1]:
        p2.add_run(',')
p2.add_run(']')
a2 = doc.add_paragraph()
a2.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a2.add_run(str(sum(myDistances_2)/(noAtoms)))

r2 = doc.add_paragraph()
r2.add_run('Radius of gyration = ').bold = True
r2.add_run(str(radiusOfGyration_2))

h3 = doc.add_paragraph('The structure number 3: ', style='Intense Quote')
d3 = doc.add_paragraph()
d3.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p3 =doc.add_paragraph('[')
for i in myDistances_3:
    p3.add_run(str(i))
    if i != myDistances_3[-1]:
        p3.add_run(',')
p3.add_run(']')
a3 = doc.add_paragraph()
a3.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a3.add_run(str(sum(myDistances_3)/(noAtoms)))

r3 = doc.add_paragraph()
r3.add_run('Radius of gyration = ').bold = True
r3.add_run(str(radiusOfGyration_3))

h4 = doc.add_paragraph('The structure number 4: ', style='Intense Quote')
d4 = doc.add_paragraph()
d4.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p4 =doc.add_paragraph('[')
for i in myDistances_4:
    p4.add_run(str(i))
    if i != myDistances_4[-1]:
        p4.add_run(',')
p4.add_run(']')
a4 = doc.add_paragraph()
a4.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a4.add_run(str(sum(myDistances_4)/(noAtoms)))

r4 = doc.add_paragraph()
r4.add_run('Radius of gyration = ').bold = True
r4.add_run(str(radiusOfGyration_4))

h5 = doc.add_paragraph('The structure number 5: ', style='Intense Quote')
d5 = doc.add_paragraph()
d5.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p5 =doc.add_paragraph('[')
for i in myDistances_5:
    p5.add_run(str(i))
    if i != myDistances_5[-1]:
        p5.add_run(',')
p5.add_run(']')
a5 = doc.add_paragraph()
a5.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a5.add_run(str(sum(myDistances_5)/(noAtoms)))

r5 = doc.add_paragraph()
r5.add_run('Radius of gyration = ').bold = True
r5.add_run(str(radiusOfGyration_5))

h6 = doc.add_paragraph('The structure number 6: ', style='Intense Quote')
d6 = doc.add_paragraph()
d6.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p6 =doc.add_paragraph('[')
for i in myDistances_6:
    p6.add_run(str(i))
    if i != myDistances_6[-1]:
        p6.add_run(',')
p6.add_run(']')
a6 = doc.add_paragraph()
a6.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a6.add_run(str(sum(myDistances_6)/(noAtoms)))

r6 = doc.add_paragraph()
r6.add_run('Radius of gyration = ').bold = True
r6.add_run(str(radiusOfGyration_6))

h7 = doc.add_paragraph('The structure number 7: ', style='Intense Quote')
d7 = doc.add_paragraph()
d7.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p7 =doc.add_paragraph('[')
for i in myDistances_7:
    p7.add_run(str(i))
    if i != myDistances_7[-1]:
        p7.add_run(',')
p7.add_run(']')
a7 = doc.add_paragraph()
a7.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a7.add_run(str(sum(myDistances_7)/(noAtoms)))

r7 = doc.add_paragraph()
r7.add_run('Radius of gyration = ').bold = True
r7.add_run(str(radiusOfGyration_7))

h8 = doc.add_paragraph('The structure number 8: ', style='Intense Quote')
d8 = doc.add_paragraph()
d8.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p8 =doc.add_paragraph('[')
for i in myDistances_8:
    p8.add_run(str(i))
    if i != myDistances_8[-1]:
        p8.add_run(',')
p8.add_run(']')
a8 = doc.add_paragraph()
a8.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a8.add_run(str(sum(myDistances_8)/(noAtoms)))

r8 = doc.add_paragraph()
r8.add_run('Radius of gyration = ').bold = True
r8.add_run(str(radiusOfGyration_8))

h9 = doc.add_paragraph('The structure number 9: ', style='Intense Quote')
d9 = doc.add_paragraph()
d9.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p9 =doc.add_paragraph('[')
for i in myDistances_9:
    p9.add_run(str(i))
    if i != myDistances_9[-1]:
        p9.add_run(',')
p9.add_run(']')
a9 = doc.add_paragraph()
a9.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a9.add_run(str(sum(myDistances_9)/(noAtoms)))

r9 = doc.add_paragraph()
r9.add_run('Radius of gyration = ').bold = True
r9.add_run(str(radiusOfGyration_9))

h10 = doc.add_paragraph('The structure number 10: ', style='Intense Quote')
d10 = doc.add_paragraph()
d10.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p10 =doc.add_paragraph('[')
for i in myDistances_10:
    p10.add_run(str(i))
    if i != myDistances_10[-1]:
        p10.add_run(',')
p10.add_run(']')
a10 = doc.add_paragraph()
a10.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a10.add_run(str(sum(myDistances_10)/(noAtoms)))

r10 = doc.add_paragraph()
r10.add_run('Radius of gyration = ').bold = True
r10.add_run(str(radiusOfGyration_10))

h11 = doc.add_paragraph('The structure number 11: ', style='Intense Quote')
d11 = doc.add_paragraph()
d11.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p11 =doc.add_paragraph('[')
for i in myDistances_11:
    p11.add_run(str(i))
    if i != myDistances_11[-1]:
        p11.add_run(',')
p11.add_run(']')
a11 = doc.add_paragraph()
a11.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a11.add_run(str(sum(myDistances_11)/(noAtoms)))

r11 = doc.add_paragraph()
r11.add_run('Radius of gyration = ').bold = True
r11.add_run(str(radiusOfGyration_11))

h12 = doc.add_paragraph('The structure number 12: ', style='Intense Quote')
d12 = doc.add_paragraph()
d12.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p12 =doc.add_paragraph('[')
for i in myDistances_12:
    p12.add_run(str(i))
    if i != myDistances_12[-1]:
        p12.add_run(',')
p12.add_run(']')
a12 = doc.add_paragraph()
a12.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a12.add_run(str(sum(myDistances_12)/(noAtoms)))

r12 = doc.add_paragraph()
r12.add_run('Radius of gyration = ').bold = True
r12.add_run(str(radiusOfGyration_12))

h13 = doc.add_paragraph('The structure number 13: ', style='Intense Quote')
d13 = doc.add_paragraph()
d13.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p13 =doc.add_paragraph('[')
for i in myDistances_13:
    p13.add_run(str(i))
    if i != myDistances_13[-1]:
        p13.add_run(',')
p13.add_run(']')
a13 = doc.add_paragraph()
a13.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a13.add_run(str(sum(myDistances_13)/(noAtoms)))

r13 = doc.add_paragraph()
r13.add_run('Radius of gyration = ').bold = True
r13.add_run(str(radiusOfGyration_13))

h14 = doc.add_paragraph('The structure number 14: ', style='Intense Quote')
d14 = doc.add_paragraph()
d14.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p14 =doc.add_paragraph('[')
for i in myDistances_14:
    p14.add_run(str(i))
    if i != myDistances_14[-1]:
        p14.add_run(',')
p14.add_run(']')
a14 = doc.add_paragraph()
a14.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a14.add_run(str(sum(myDistances_14)/(noAtoms)))

r14 = doc.add_paragraph()
r14.add_run('Radius of gyration = ').bold = True
r14.add_run(str(radiusOfGyration_14))

h15 = doc.add_paragraph('The structure number 15: ', style='Intense Quote')
d15 = doc.add_paragraph()
d15.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p15 =doc.add_paragraph('[')
for i in myDistances_15:
    p15.add_run(str(i))
    if i != myDistances_15[-1]:
        p15.add_run(',')
p15.add_run(']')
a15 = doc.add_paragraph()
a15.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a15.add_run(str(sum(myDistances_15)/(noAtoms)))

r15 = doc.add_paragraph()
r15.add_run('Radius of gyration = ').bold = True
r15.add_run(str(radiusOfGyration_15))

h16 = doc.add_paragraph('The structure number 16: ', style='Intense Quote')
d16 = doc.add_paragraph()
d16.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p16 =doc.add_paragraph('[')
for i in myDistances_16:
    p16.add_run(str(i))
    if i != myDistances_16[-1]:
        p16.add_run(',')
p16.add_run(']')
a16 = doc.add_paragraph()
a16.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a16.add_run(str(sum(myDistances_16)/(noAtoms)))

r16 = doc.add_paragraph()
r16.add_run('Radius of gyration = ').bold = True
r16.add_run(str(radiusOfGyration_16))

h17 = doc.add_paragraph('The structure number 17: ', style='Intense Quote')
d17 = doc.add_paragraph()
d17.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p17 =doc.add_paragraph('[')
for i in myDistances_17:
    p17.add_run(str(i))
    if i != myDistances_17[-1]:
        p17.add_run(',')
p17.add_run(']')
a17 = doc.add_paragraph()
a17.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a17.add_run(str(sum(myDistances_17)/(noAtoms)))

r17 = doc.add_paragraph()
r17.add_run('Radius of gyration = ').bold = True
r17.add_run(str(radiusOfGyration_17))

h18 = doc.add_paragraph('The structure number 18: ', style='Intense Quote')
d18 = doc.add_paragraph()
d18.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p18 =doc.add_paragraph('[')
for i in myDistances_18:
    p18.add_run(str(i))
    if i != myDistances_18[-1]:
        p18.add_run(',')
p18.add_run(']')
a18 = doc.add_paragraph()
a18.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a18.add_run(str(sum(myDistances_18)/(noAtoms)))

r18 = doc.add_paragraph()
r18.add_run('Radius of gyration = ').bold = True
r18.add_run(str(radiusOfGyration_18))

h19 = doc.add_paragraph('The structure number 19: ', style='Intense Quote')
d19 = doc.add_paragraph()
d19.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p19 =doc.add_paragraph('[')
for i in myDistances_19:
    p19.add_run(str(i))
    if i != myDistances_19[-1]:
        p19.add_run(',')
p19.add_run(']')
a19 = doc.add_paragraph()
a19.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a19.add_run(str(sum(myDistances_19)/(noAtoms)))

r19 = doc.add_paragraph()
r19.add_run('Radius of gyration = ').bold = True
r19.add_run(str(radiusOfGyration_19))

h20 = doc.add_paragraph('The structure number 20: ', style='Intense Quote')
d20 = doc.add_paragraph()
d20.add_run('The atom-center of mass distances of every atom in the chain = ').bold = True
p20 =doc.add_paragraph('[')
for i in myDistances_20:
    p20.add_run(str(i))
    if i != myDistances_20[-1]:
        p20.add_run(',')
p20.add_run(']')
a20 = doc.add_paragraph()
a20.add_run('Average atom-center of mass per all atoms in  the   structures = ').bold = True
a20.add_run(str(sum(myDistances_20)/(noAtoms)))

r20 = doc.add_paragraph()
r20.add_run('Radius of gyration = ').bold = True
r20.add_run(str(radiusOfGyration_20))



introb11 = doc.add_paragraph()
introb11.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb11.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb11.add_run('Data for the last ten structures:-').bold = True
hb11 = doc.add_paragraph('structure number 11: ', style='Intense Quote')
db11 = doc.add_paragraph()
db11.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb11 = doc.add_paragraph('[')
for i in Densities_11:
    pb11.add_run(str(i))
    if i != Densities_11[-1]:
        pb11.add_run(',')
pb11.add_run(']')
ab11 = doc.add_paragraph()
ab11.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab11.add_run('[')
for i in volumes_11:
    ab11.add_run(str(i))
    if i != volumes_11[-1]:
        ab11.add_run(',')
ab11.add_run(']')
rb11 = doc.add_paragraph()
rb11.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb11.add_run('[')
for i in B_11:
    rb11.add_run(str(i))
    if i != B_11[-1]:
        rb11.add_run(',')
rb11.add_run(']')

introb12 = doc.add_paragraph()
introb12.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb12.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb12.add_run('Data for the last ten structures:-').bold = True
hb12 = doc.add_paragraph('structure number 12: ', style='Intense Quote')
db12 = doc.add_paragraph()
db12.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb12 = doc.add_paragraph('[')
for i in Densities_12:
    pb12.add_run(str(i))
    if i != Densities_12[-1]:
        pb12.add_run(',')
pb12.add_run(']')
ab12 = doc.add_paragraph()
ab12.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab12.add_run('[')
for i in volumes_12:
    ab12.add_run(str(i))
    if i != volumes_12[-1]:
        ab12.add_run(',')
ab12.add_run(']')
rb12 = doc.add_paragraph()
rb12.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb12.add_run('[')
for i in B_12:
    rb12.add_run(str(i))
    if i != B_12[-1]:
        rb12.add_run(',')
rb12.add_run(']')

introb13 = doc.add_paragraph()
introb13.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb13.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb13.add_run('Data for the last ten structures:-').bold = True
hb13 = doc.add_paragraph('structure number 13: ', style='Intense Quote')
db13 = doc.add_paragraph()
db13.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb13 = doc.add_paragraph('[')
for i in Densities_13:
    pb13.add_run(str(i))
    if i != Densities_13[-1]:
        pb13.add_run(',')
pb13.add_run(']')
ab13 = doc.add_paragraph()
ab13.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab13.add_run('[')
for i in volumes_13:
    ab13.add_run(str(i))
    if i != volumes_13[-1]:
        ab13.add_run(',')
ab13.add_run(']')
rb13 = doc.add_paragraph()
rb13.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb13.add_run('[')
for i in B_13:
    rb13.add_run(str(i))
    if i != B_13[-1]:
        rb13.add_run(',')
rb13.add_run(']')

introb14 = doc.add_paragraph()
introb14.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb14.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb14.add_run('Data for the last ten structures:-').bold = True
hb14 = doc.add_paragraph('structure number 14: ', style='Intense Quote')
db14 = doc.add_paragraph()
db14.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb14 = doc.add_paragraph('[')
for i in Densities_14:
    pb14.add_run(str(i))
    if i != Densities_14[-1]:
        pb14.add_run(',')
pb14.add_run(']')
ab14 = doc.add_paragraph()
ab14.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab14.add_run('[')
for i in volumes_14:
    ab14.add_run(str(i))
    if i != volumes_14[-1]:
        ab14.add_run(',')
ab14.add_run(']')
rb14 = doc.add_paragraph()
rb14.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb14.add_run('[')
for i in B_14:
    rb14.add_run(str(i))
    if i != B_14[-1]:
        rb14.add_run(',')
rb14.add_run(']')

introb15 = doc.add_paragraph()
introb15.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb15.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb15.add_run('Data for the last ten structures:-').bold = True
hb15 = doc.add_paragraph('structure number 15: ', style='Intense Quote')
db15 = doc.add_paragraph()
db15.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb15 = doc.add_paragraph('[')
for i in Densities_15:
    pb15.add_run(str(i))
    if i != Densities_15[-1]:
        pb15.add_run(',')
pb15.add_run(']')
ab15 = doc.add_paragraph()
ab15.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab15.add_run('[')
for i in volumes_15:
    ab15.add_run(str(i))
    if i != volumes_15[-1]:
        ab15.add_run(',')
ab15.add_run(']')
rb15 = doc.add_paragraph()
rb15.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb15.add_run('[')
for i in B_15:
    rb15.add_run(str(i))
    if i != B_15[-1]:
        rb15.add_run(',')
rb15.add_run(']')

introb16 = doc.add_paragraph()
introb16.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb16.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb16.add_run('Data for the last ten structures:-').bold = True
hb16 = doc.add_paragraph('structure number 16: ', style='Intense Quote')
db16 = doc.add_paragraph()
db16.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb16 = doc.add_paragraph('[')
for i in Densities_16:
    pb16.add_run(str(i))
    if i != Densities_16[-1]:
        pb16.add_run(',')
pb16.add_run(']')
ab16 = doc.add_paragraph()
ab16.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab16.add_run('[')
for i in volumes_16:
    ab16.add_run(str(i))
    if i != volumes_16[-1]:
        ab16.add_run(',')
ab16.add_run(']')
rb16 = doc.add_paragraph()
rb16.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb16.add_run('[')
for i in B_16:
    rb16.add_run(str(i))
    if i != B_16[-1]:
        rb16.add_run(',')
rb16.add_run(']')

introb17 = doc.add_paragraph()
introb17.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb17.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb17.add_run('Data for the last ten structures:-').bold = True
hb17 = doc.add_paragraph('structure number 17: ', style='Intense Quote')
db17 = doc.add_paragraph()
db17.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb17 = doc.add_paragraph('[')
for i in Densities_17:
    pb17.add_run(str(i))
    if i != Densities_17[-1]:
        pb17.add_run(',')
pb17.add_run(']')
ab17 = doc.add_paragraph()
ab17.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab17.add_run('[')
for i in volumes_17:
    ab17.add_run(str(i))
    if i != volumes_17[-1]:
        ab17.add_run(',')
ab17.add_run(']')
rb17 = doc.add_paragraph()
rb17.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb17.add_run('[')
for i in B_17:
    rb17.add_run(str(i))
    if i != B_17[-1]:
        rb17.add_run(',')
rb17.add_run(']')

introb18 = doc.add_paragraph()
introb18.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb18.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb18.add_run('Data for the last ten structures:-').bold = True
hb18 = doc.add_paragraph('structure number 18: ', style='Intense Quote')
db18 = doc.add_paragraph()
db18.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb18 = doc.add_paragraph('[')
for i in Densities_18:
    pb18.add_run(str(i))
    if i != Densities_18[-1]:
        pb18.add_run(',')
pb18.add_run(']')
ab18 = doc.add_paragraph()
ab18.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab18.add_run('[')
for i in volumes_18:
    ab18.add_run(str(i))
    if i != volumes_18[-1]:
        ab18.add_run(',')
ab18.add_run(']')
rb18 = doc.add_paragraph()
rb18.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb18.add_run('[')
for i in B_18:
    rb18.add_run(str(i))
    if i != B_18[-1]:
        rb18.add_run(',')
rb18.add_run(']')

introb19 = doc.add_paragraph()
introb19.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb19.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb19.add_run('Data for the last ten structures:-').bold = True
hb19 = doc.add_paragraph('structure number 19: ', style='Intense Quote')
db19 = doc.add_paragraph()
db19.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb19 = doc.add_paragraph('[')
for i in Densities_19:
    pb19.add_run(str(i))
    if i != Densities_19[-1]:
        pb19.add_run(',')
pb19.add_run(']')
ab19 = doc.add_paragraph()
ab19.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab19.add_run('[')
for i in volumes_19:
    ab19.add_run(str(i))
    if i != volumes_19[-1]:
        ab19.add_run(',')
ab19.add_run(']')
rb19 = doc.add_paragraph()
rb19.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb19.add_run('[')
for i in B_19:
    rb19.add_run(str(i))
    if i != B_19[-1]:
        rb19.add_run(',')
rb19.add_run(']')

introb20 = doc.add_paragraph()
introb20.add_run('The data taken from the histogram plots are:'+'\n').bold = True
introb20.add_run('A) density, number of atoms and volumes of whole structure.'+'\n'+'B) density, number of atoms and volumes of individual shell.'+'\n').italic = True
introb20.add_run('Data for the last ten structures:-').bold = True
hb20 = doc.add_paragraph('structure number 20: ', style='Intense Quote')
db20 = doc.add_paragraph()
db20.add_run('Density of the shells in the whole structure respectively = ').bold = True
pb20 = doc.add_paragraph('[')
for i in Densities_20:
    pb20.add_run(str(i))
    if i != Densities_20[-1]:
        pb20.add_run(',')
pb20.add_run(']')
ab20 = doc.add_paragraph()
ab20.add_run('Volumes of the shells in the structure(width"cut from x-axis"*hight"number of atoms")  respectively = '+'\n').bold =True
ab20.add_run('[')
for i in volumes_20:
    ab20.add_run(str(i))
    if i != volumes_20[-1]:
        ab20.add_run(',')
ab20.add_run(']')
rb20 = doc.add_paragraph()
rb20.add_run('Number of atoms in the shells of the structure respectively = '+'\n').bold = True
rb20.add_run('[')
for i in B_20:
    rb20.add_run(str(i))
    if i != B_20[-1]:
        rb20.add_run(',')
rb20.add_run(']')


doc.save('Discussion_30M.docx')
print("The number of atoms in the shell number 1 in the last ten structures respectively =", [L_11[0],L_12[0],L_13[0],L_14[0],L_15[0],L_16[0],L_17[0],L_18[0],L_19[0],L_20[0]])
print("The number of atoms in the shell number 2 in the last ten structures respectively =", [L_11[1],L_12[1],L_13[1],L_14[1],L_15[1],L_16[1],L_17[1],L_18[1],L_19[1],L_20[1]])
print("The number of atoms in the shell number 3 in the last ten structures respectively =", [L_11[2],L_12[2],L_13[2],L_14[2],L_15[2],L_16[2],L_17[2],L_18[2],L_19[2],L_20[2]])
print("The number of atoms in the shell number 4 in the last ten structures respectively =", [L_11[3],L_12[3],L_13[3],L_14[3],L_15[3],L_16[3],L_17[3],L_18[3],L_19[3],L_20[3]])
print("The number of atoms in the shell number 5 in the last ten structures respectively =", [L_11[4],L_12[4],L_13[4],L_14[4],L_15[4],L_16[4],L_17[4],L_18[4],L_19[4],L_20[4]])
print("The number of atoms in the shell number 6 in the last ten structures respectively =", [L_11[5],L_12[5],L_13[5],L_14[5],L_15[5],L_16[5],L_17[5],L_18[5],L_19[5],L_20[5]])
print("The number of atoms in the shell number 7 in the last ten structures respectively =", [L_11[6],L_12[6],L_13[6],L_14[6],L_15[6],L_16[6],L_17[6],L_18[6],L_19[6],L_20[6]])
print("The number of atoms in the shell number 8 in the last ten structures respectively =", [L_11[7],L_12[7],L_13[7],L_14[7],L_15[7],L_16[7],L_17[7],L_18[7],L_19[7],L_20[7]])
print("The number of atoms in the shell number 9 in the last ten structures respectively =", [L_11[8],L_12[8],L_13[8],L_14[8],L_15[8],L_16[8],L_17[8],L_18[8],L_19[8],L_20[8]])
print("The number of atoms in the shell number 10 in the last ten structures respectively =", [L_11[9],L_12[9],L_13[9],L_14[9],L_15[9],L_16[9],L_17[9],L_18[9],L_19[9],L_20[9]])
